from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Styles helper ────────────────────────────────────────────────────────────
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x53, 0xff)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1f, 0x3a, 0x93)

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x21, 0x74, 0xb5)

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x27, 0x6b, 0x29)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)

def add_qa(doc, q_num, question, answer):
    p = doc.add_paragraph()
    r = p.add_run(f"Q{q_num}. {question}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1a, 0x53, 0xff)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Answer: ")
    r2.bold = True
    r2.font.size = Pt(11)
    p2.add_run(answer).font.size = Pt(11)
    doc.add_paragraph()

def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
add_title(doc, "LOAN MAKER PROJECT")
doc.add_paragraph()
p = doc.add_paragraph("Complete Technical Documentation & Interview Preparation Guide")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p.runs[0].font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph("Stack: Spring Boot 3.5 | React 19 | Python FastAPI | MySQL | JWT | AI/ML")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "1. PROJECT OVERVIEW")
add_body(doc, "Loan Maker is a full-stack, AI-powered loan management web application that allows users to register, submit assets, apply for loans from various bank providers, and receive instant AI-driven loan decisions. Administrators can manage users, view all loan applications, approve or reject them, and promote users to admin role.")

doc.add_paragraph()
add_h2(doc, "1.1 Key Features")
feats = [
    "User Registration & Login with JWT-based authentication",
    "Role-based access control: USER and ADMIN roles",
    "Asset Management: users can add/delete assets (Gold, Land, Property)",
    "Loan Application: apply with a selected bank provider and asset as collateral",
    "AI Auto-Evaluation: rule-based + ML model decides APPROVED / PENDING / REJECTED instantly",
    "AI Eligibility Test: Python FastAPI microservice checks eligibility and recommends lenders",
    "Admin Dashboard: view all loans, approve/reject, view statistics, promote/demote users",
    "Loan History: users can track all their past loan applications",
    "Live Deployed: Frontend on Netlify, Backend on Railway, AI on Render, DB on Aiven (MySQL Cloud)",
]
for f in feats:
    add_bullet(doc, f)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "2. TECHNOLOGY STACK (A to Z)")

add_h2(doc, "2.1 Backend — Spring Boot (Java)")
stack_back = [
    ("Spring Boot 3.5.5 (Java 21): ", "Core framework providing REST API, auto-configuration, embedded Tomcat server."),
    ("Spring Web (spring-boot-starter-web): ", "Enables @RestController, @RequestMapping, and JSON responses via Jackson."),
    ("Spring Data JPA (spring-boot-starter-data-jpa): ", "ORM layer; maps Java entities to MySQL tables using Hibernate under the hood. Uses CrudRepository/JpaRepository for CRUD operations."),
    ("Spring Security (spring-boot-starter-security): ", "Provides authentication, authorization, password hashing, CSRF management, CORS configuration, and stateless session handling."),
    ("JWT — JJWT 0.11.5 (jjwt-api, jjwt-impl, jjwt-jackson): ", "JSON Web Token generation and validation. Tokens signed with HS256 algorithm, expire in 1 hour. Stored in browser localStorage."),
    ("BCryptPasswordEncoder: ", "Secure password hashing with adaptive cost factor. Passwords are never stored in plain text."),
    ("Lombok: ", "Reduces boilerplate with @Getter, @Setter, @NoArgsConstructor, @AllArgsConstructor, @Builder, @Data annotations."),
    ("MySQL Connector (mysql-connector-j): ", "JDBC driver for connecting Spring Boot to MySQL."),
    ("Spring Boot Validation (spring-boot-starter-validation): ", "Enables @Valid and constraint annotations for request body validation."),
    ("RestTemplate: ", "Used internally to call the Python AI microservice from Spring Boot."),
    ("Maven: ", "Build tool; pom.xml manages all dependencies, plugins, and build lifecycle."),
    ("Dockerfile: ", "Containerizes the backend; deployed via Railway cloud."),
]
for k, v in stack_back:
    add_bullet(doc, v, bold_prefix=k)

add_h2(doc, "2.2 Database — MySQL (Cloud: Aiven)")
db_points = [
    "Hosted on Aiven Cloud (managed MySQL 8 service) with SSL connection.",
    "spring.jpa.hibernate.ddl-auto=update — Hibernate auto-creates/updates tables on startup.",
    "4 tables: users, loan_applications, loan_providers, assets.",
    "Relational: LoanApplication has ManyToOne to User, LoanProvider, and Asset.",
]
for d in db_points:
    add_bullet(doc, d)

add_h2(doc, "2.3 Frontend — React 19 + Vite")
stack_front = [
    ("React 19: ", "Latest React with hooks (useState, useEffect). Functional component architecture."),
    ("Vite 7: ", "Next-gen frontend build tool. Much faster than Create React App. Hot module replacement (HMR) in dev."),
    ("React Router DOM v7: ", "Client-side routing. ProtectedRoute component guards authenticated routes. Supports role-based routing."),
    ("Tailwind CSS v4: ", "Utility-first CSS framework. Enables rapid UI building with responsive design."),
    ("Axios: ", "HTTP client for making API calls to the backend. Configured with base URL and auth headers."),
    ("Recharts: ", "Responsive charting library for admin dashboard statistics (bar charts, pie charts)."),
    ("Framer Motion: ", "Animation library for smooth page transitions and UI animations."),
    ("Lucide React: ", "Icon library providing consistent SVG icons (CheckCircle, XCircle, DollarSign, etc.)."),
    ("react-hot-toast & react-toastify: ", "Toast notification systems for user feedback messages."),
    ("jwt-decode: ", "Decodes JWT on the frontend to extract user role/identity without calling the server."),
    ("react-csv: ", "Enables CSV export of loan data in admin reports."),
    ("localStorage: ", "Stores JWT token, cached eligibility results, and last-notified loan status for UX continuity."),
]
for k, v in stack_front:
    add_bullet(doc, v, bold_prefix=k)

add_h2(doc, "2.4 AI Microservice — Python FastAPI")
stack_ai = [
    ("FastAPI: ", "Modern, high-performance Python web framework. Auto-generates OpenAPI/Swagger docs."),
    ("scikit-learn: ", "Machine learning library. A pre-trained model (loan_model.pkl) is loaded using pickle to predict loan approval (binary classification: 0 or 1)."),
    ("Pydantic: ", "Request/response data validation. LoanRequest model enforces types (age, income, credit_score, etc.)."),
    ("NumPy: ", "Used to construct feature arrays from request data for the ML model."),
    ("Uvicorn: ", "ASGI server running the FastAPI app in production."),
    ("Render.com: ", "Cloud platform where the AI microservice is deployed (https://loan-maker-ai.onrender.com)."),
]
for k, v in stack_ai:
    add_bullet(doc, v, bold_prefix=k)

add_h2(doc, "2.5 Deployment Architecture")
add_bullet(doc, "Frontend → Netlify (https://timely-tanuki-loan-maker.netlify.app)")
add_bullet(doc, "Backend → Railway (https://loan-maker-backend-production.up.railway.app)")
add_bullet(doc, "AI Microservice → Render (https://loan-maker-ai.onrender.com)")
add_bullet(doc, "Database → Aiven Cloud MySQL (mysql-31120da3-loan-maker-db.i.aivencloud.com:25875)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE & FLOW
# ═══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "3. SYSTEM ARCHITECTURE & FLOW")

add_h2(doc, "3.1 High-Level Architecture")
add_body(doc, "The system follows a 3-tier architecture with an additional AI microservice:")
add_bullet(doc, "Tier 1 — Presentation Layer: React frontend (SPA) served from Netlify CDN")
add_bullet(doc, "Tier 2 — Application Layer: Spring Boot REST API on Railway with JWT security")
add_bullet(doc, "Tier 2.5 — AI Layer: Python FastAPI microservice on Render (called by Spring Boot)")
add_bullet(doc, "Tier 3 — Data Layer: MySQL on Aiven Cloud (persistent storage)")

add_h2(doc, "3.2 Authentication Flow")
steps = [
    "User submits login form (email + password) → POST /api/auth/login",
    "AuthController calls AuthenticationManager.authenticate() with credentials",
    "UserDetailsServiceImpl loads user by email from database",
    "BCryptPasswordEncoder verifies hashed password",
    "On success, JwtUtil.generateToken() creates a signed JWT (HS256, 1-hour expiry)",
    "JWT returned to frontend, stored in localStorage",
    "All subsequent requests include JWT in Authorization header (Bearer token)",
    "JwtFilter intercepts each request, validates token, and sets SecurityContext",
]
for i, s in enumerate(steps, 1):
    add_bullet(doc, f"Step {i}: {s}")

add_h2(doc, "3.3 Loan Application Flow")
loan_steps = [
    "User adds asset (POST /api/assets/{userId}) — asset linked to user account",
    "User visits Offers page — POST /api/ai/predict-and-recommend sends profile to AI microservice",
    "AI microservice returns eligibility score + recommended bank providers",
    "User selects a bank and applies — POST /api/loans/apply/{userId}/{providerId}/{assetId}",
    "LoanService.evaluateLoan() runs rule-based check (credit score, previous loans, asset value)",
    "Status is APPROVED/PENDING/REJECTED — saved to loan_applications table",
    "User dashboard shows notification modal for latest status change",
    "Admin can also auto-evaluate via PUT /api/loans/admin/{loanId}/auto",
    "Admin can manually approve (PUT /api/loans/admin/{loanId}/approve) or reject",
]
for i, s in enumerate(loan_steps, 1):
    add_bullet(doc, f"Step {i}: {s}")

add_h2(doc, "3.4 AI Evaluation Logic")
add_body(doc, "Two mechanisms work together:")
add_bullet(doc, "Rule-Based (Spring Boot - LoanService.evaluateLoan()): If creditScore >= 700 AND previousLoans <= 2 AND loanAmount <= assetValue → APPROVED. If creditScore < 500 OR loanAmount > assetValue * 1.2 → REJECTED. Otherwise → PENDING.")
add_bullet(doc, "ML Model (Python FastAPI): Trained scikit-learn model (loan_model.pkl) uses features: age, income, existing_emi, credit_score, asset_value, loan_amount, tenure. Returns boolean: approved (true/false).")
add_bullet(doc, "Recommendation Engine (Python - loan_recommendation.py): Returns a ranked list of lenders suitable for the user's profile.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. DATABASE DESIGN
# ═══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "4. DATABASE DESIGN")

add_h2(doc, "4.1 Entity-Relationship Summary")

add_h3(doc, "Table: users")
for col in ["id (PK, AUTO_INCREMENT)", "name", "email (unique)", "password (BCrypt hash)", "phone", "role (USER / ROLE_ADMIN)", "previousLoans (int)", "creditScore (int, default 700)"]:
    add_bullet(doc, col)

add_h3(doc, "Table: loan_providers")
for col in ["id (PK)", "bankName", "interestRate (double)", "maxAmount (double)", "minCreditScore (int)", "loanType (Personal/Gold/Land/etc.)"]:
    add_bullet(doc, col)

add_h3(doc, "Table: assets")
for col in ["id (PK)", "assetType (Gold/Land/Property)", "description", "value (double)", "user_id (FK → users.id)"]:
    add_bullet(doc, col)

add_h3(doc, "Table: loan_applications")
for col in ["id (PK)", "loanAmount (double)", "tenure (int, months)", "status (PENDING/APPROVED/REJECTED)", "user_id (FK → users.id)", "provider_id (FK → loan_providers.id)", "asset_id (FK → assets.id)"]:
    add_bullet(doc, col)

add_h2(doc, "4.2 Relationships")
add_bullet(doc, "User → Assets: One-to-Many (one user can have many assets)")
add_bullet(doc, "User → LoanApplications: One-to-Many (one user can apply for many loans)")
add_bullet(doc, "LoanApplication → LoanProvider: Many-to-One")
add_bullet(doc, "LoanApplication → Asset: Many-to-One (asset used as collateral)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. BACKEND COMPONENT DETAILS
# ═══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "5. BACKEND COMPONENT DETAILS")

add_h2(doc, "5.1 Controllers (REST API Endpoints)")

add_h3(doc, "AuthController — /api/auth")
add_bullet(doc, "POST /login — authenticates user, returns JWT token")
add_bullet(doc, "POST /register — creates new user with BCrypt password, returns JWT + user details")

add_h3(doc, "LoanController — /api/loans")
add_bullet(doc, "GET /providers/{loanType} — fetch providers by loan type")
add_bullet(doc, "POST /apply/{userId}/{providerId}/{assetId} — apply for loan (auto AI evaluation)")
add_bullet(doc, "GET /applications/{userId} — get logged-in user's loans")
add_bullet(doc, "GET /admin/all — all loans (ADMIN only)")
add_bullet(doc, "GET /admin/pending — pending loans (ADMIN only)")
add_bullet(doc, "PUT /admin/{loanId}/approve — approve loan (ADMIN only)")
add_bullet(doc, "PUT /admin/{loanId}/reject — reject loan (ADMIN only)")
add_bullet(doc, "PUT /admin/{loanId}/auto — auto-evaluate using AI rules (ADMIN only)")
add_bullet(doc, "GET /stats — loan statistics (total/approved/pending/rejected)")

add_h3(doc, "AssetController — /api/assets")
add_bullet(doc, "POST /{userId} — add asset for a user")
add_bullet(doc, "GET /{userId} — get all assets of a user")
add_bullet(doc, "DELETE /{assetId} — delete an asset")

add_h3(doc, "AdminController — /api/admin")
add_bullet(doc, "GET /users — list all users (ADMIN only)")
add_bullet(doc, "PUT /users/promote/{userId} — promote user to ADMIN")
add_bullet(doc, "PUT /users/demote/{userId} — demote admin to USER")

add_h3(doc, "LoanAIController — /api/ai")
add_bullet(doc, "POST /predict — sends data to Python AI, returns approval boolean")
add_bullet(doc, "POST /recommend — gets recommended lenders from Python AI")
add_bullet(doc, "POST /predict-and-recommend — combined call returning both prediction and recommendations")

add_h2(doc, "5.2 Security Layer")
add_bullet(doc, "JwtFilter: extends OncePerRequestFilter. Extracts Bearer token from Authorization header, validates it, loads UserDetails, and sets SecurityContext.")
add_bullet(doc, "JwtUtil: generates/validates tokens using HMAC-SHA256 (HS256). Secret from environment variable. 1-hour expiry.")
add_bullet(doc, "SecurityConfig: CSRF disabled (stateless REST). CORS configured for localhost:5173 and Netlify URL. Public endpoints: /api/auth/**, /api/users/register, /api/ai/**. All others require authentication.")
add_bullet(doc, "UserDetailsServiceImpl: loads user by email, returns Spring Security UserDetails with GrantedAuthority based on role field.")
add_bullet(doc, "@PreAuthorize('hasAuthority(\"ROLE_ADMIN\")'): method-level security on admin endpoints.")

add_h2(doc, "5.3 Service Layer")
add_bullet(doc, "LoanService: business logic for applying, evaluating, approving, rejecting loans. Contains AI rule engine.")
add_bullet(doc, "LoanAIService: uses RestTemplate to call Python FastAPI microservice at loan-maker-ai.onrender.com.")
add_bullet(doc, "UserService: get all users, promote/demote roles.")
add_bullet(doc, "AssetService: save and retrieve assets by user.")
add_bullet(doc, "LoanRecommendationService: calls /recommend endpoint of Python AI.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. FRONTEND COMPONENT DETAILS
# ═══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "6. FRONTEND COMPONENT DETAILS")

add_h2(doc, "6.1 Page Structure")
pages = [
    ("Auth Pages", "Login.jsx — email/password form, stores JWT in localStorage. Register.jsx — new user registration form."),
    ("User Pages", "Dashboard.jsx — shows profile, loan offers, recommended loans, and loan status modal notification. LoanApply.jsx — select loan type, provider, asset, enter amount, submit application. LoanHistory.jsx — displays all past loan applications with status badges. ManageAssets.jsx — add/delete user assets. Offers.jsx — AI eligibility test form, displays score and recommended banks. Profile.jsx — view and update profile details."),
    ("Admin Pages", "Dashboard.jsx (Admin) — statistics cards, charts (Recharts). LoanList.jsx — all loan applications with approve/reject/auto-evaluate buttons. UserList.jsx — all users with promote/demote buttons. Reports.jsx — CSV export of loan data. AddAdmin.jsx — creates new admin account."),
]
for name, desc in pages:
    add_bullet(doc, desc, bold_prefix=f"{name}: ")

add_h2(doc, "6.2 Routing (React Router DOM v7)")
add_bullet(doc, "AppRoutes.jsx defines all application routes")
add_bullet(doc, "ProtectedRoute.jsx: checks localStorage for token; if not present redirects to /login")
add_bullet(doc, "Role-based routing: ADMIN users see admin dashboard; USER users see user dashboard")
add_bullet(doc, "jwt-decode decodes the stored JWT on the frontend to extract role without a server call")

add_h2(doc, "6.3 State Management")
add_body(doc, "The project uses React's built-in useState and useEffect hooks for local component state. No Redux or Context API — the JWT token and user ID stored in localStorage serve as the global state source. Eligibility results are cached in localStorage under key eligibility_result_{userId}.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. INTERVIEW QUESTIONS & ANSWERS
# ═══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "7. INTERVIEW QUESTIONS & ANSWERS")

add_h2(doc, "Section A — Project Overview Questions")

add_qa(doc, 1,
       "Can you briefly explain the Loan Maker project?",
       "Loan Maker is a full-stack AI-powered loan management system. The frontend is built with React 19 and Vite, the backend with Spring Boot 3.5 and Java 21, and there's a separate Python FastAPI AI microservice. Users can register, add assets as collateral, apply for loans from different bank providers, and receive instant AI-driven loan decisions. Admins can manage users, approve/reject loans, and view statistics. The whole system is deployed on cloud platforms — Netlify for frontend, Railway for the backend, Render for AI, and Aiven for MySQL.")

add_qa(doc, 2,
       "What was your role in this project?",
       "I designed and implemented the complete system — backend REST API with Spring Boot, JWT security layer, database design, frontend React components, and integrated the Python AI microservice. I also configured the CORS policies and deployed each layer to its respective cloud platform.")

add_qa(doc, 3,
       "Why did you choose this technology stack?",
       "Spring Boot provides rapid REST API development with production-grade features. React with Vite gives a fast, modern frontend experience. MySQL is a reliable relational database suitable for financial data. Python with FastAPI was chosen for the AI microservice because scikit-learn and the ML ecosystem in Python are industry standard. JWT was chosen for stateless, scalable authentication suitable for a cloud-deployed API.")

add_h2(doc, "Section B — Spring Boot / Backend Questions")

add_qa(doc, 4,
       "How does JWT authentication work in your project?",
       "When a user logs in, the AuthController authenticates credentials using Spring Security's AuthenticationManager. On success, JwtUtil generates a JWT signed with HMAC-SHA256 (HS256) using a secret key. The token contains the user's email as subject and expires in 1 hour. The token is returned to the frontend and stored in localStorage. For every subsequent request, the JwtFilter extracts the Bearer token from the Authorization header, validates it, loads UserDetails, and sets the SecurityContext. This makes the system completely stateless — no server-side sessions.")

add_qa(doc, 5,
       "What is the difference between @PreAuthorize and requestMatchers in your SecurityConfig?",
       "requestMatchers in SecurityConfig handles URL-level authorization at the filter chain level — it decides whether a request is allowed to reach any controller at all. @PreAuthorize is method-level security that runs after the request reaches the controller, checking the authority of the authenticated user. I use both: requestMatchers for broad rules (e.g., permit /api/auth/**) and @PreAuthorize('hasAuthority(\"ROLE_ADMIN\")') for fine-grained endpoint protection.")

add_qa(doc, 6,
       "Why is CSRF disabled in your application?",
       "CSRF (Cross-Site Request Forgery) protection is designed for browser sessions using cookies. Since this application uses JWT stored in localStorage and sent in an Authorization header (not a cookie), CSRF attacks are not applicable. Disabling CSRF is standard practice for stateless REST APIs secured with JWT.")

add_qa(doc, 7,
       "Explain Spring Data JPA and how you used it.",
       "Spring Data JPA is an abstraction layer on top of JPA/Hibernate. By extending JpaRepository or CrudRepository, we get CRUD operations without writing SQL. Entity classes annotated with @Entity map to database tables. Relationships like @ManyToOne with @JoinColumn handle foreign keys. I used it for UserRepository.findByEmail(), LoanApplicationRepository.findByUser(), LoanApplicationRepository.findByStatus(), and LoanProviderRepository.findByLoanType() — all generated from method names by Spring Data.")

add_qa(doc, 8,
       "What is Lombok and why did you use it?",
       "Lombok is an annotation processor that generates boilerplate Java code at compile time. @Getter and @Setter generate getters/setters. @NoArgsConstructor creates a no-arg constructor. @AllArgsConstructor creates a constructor with all fields. @Builder enables the builder pattern. @Data combines @Getter, @Setter, @ToString, @EqualsAndHashCode. It reduces entity class code from ~100 lines to ~20 lines, improving readability and maintainability.")

add_qa(doc, 9,
       "How does your AI evaluation logic work in LoanService?",
       "LoanService.evaluateLoan() applies rule-based logic: if credit score >= 700, previous loans <= 2, and loan amount <= asset value, the loan is APPROVED instantly. If credit score < 500 or loan amount > 120% of asset value, it's REJECTED. Borderline cases remain PENDING for admin review. This mirrors real-world bank decision criteria. Additionally, the Python AI microservice uses a trained ML model (RandomForest/similar) for a probabilistic prediction based on 7 features including income and EMI.")

add_qa(doc, 10,
       "How do you call the Python AI microservice from Spring Boot?",
       "I use Spring's RestTemplate in LoanAIService. It sends an HTTP POST request with a LoanAiRequest DTO (serialized as JSON) to the Python FastAPI endpoints at https://loan-maker-ai.onrender.com/predict and /recommend. The response is mapped to a Map<String, Object>. Error handling wraps HttpClientErrorException to return meaningful error messages instead of crashing.")

add_qa(doc, 11,
       "What is @CrossOrigin and why did you use it?",
       "@CrossOrigin is a Spring annotation that enables CORS for specific controllers. I also configured CORS globally in SecurityConfig using CorsConfigurationSource, allowing requests from localhost:5173 (development) and the Netlify production URL. CORS must be configured so the browser doesn't block frontend requests to the backend running on a different domain/port.")

add_qa(doc, 12,
       "Explain the DTO pattern you used.",
       "LoanApplicationDTO is a Data Transfer Object that flattens the nested LoanApplication entity for the API response. Entities have JPA relationships (@ManyToOne), which can cause serialization issues or expose unnecessary data. The DTO contains only the required fields: id, userId, providerId, bankName, assetId, assetType, loanAmount, tenure, status. The toDTO() helper method in LoanController converts entities to DTOs before sending the response.")

add_h2(doc, "Section C — Spring Security / JWT Deep Dive")

add_qa(doc, 13,
       "What is BCrypt and why is it used for password storage?",
       "BCrypt is an adaptive password hashing algorithm. Unlike MD5 or SHA, BCrypt is intentionally slow and includes a 'cost factor' that can be increased as hardware improves. It also automatically generates and embeds a random salt, making rainbow table attacks useless. Even if two users have the same password, their BCrypt hashes will be different. Spring Security's BCryptPasswordEncoder handles both encoding and matching.")

add_qa(doc, 14,
       "What is the difference between ROLE_ADMIN and ADMIN in Spring Security?",
       "Spring Security has two methods: hasRole() automatically prepends 'ROLE_' to the given string, so hasRole('ADMIN') checks for 'ROLE_ADMIN'. hasAuthority() checks the exact string, so hasAuthority('ROLE_ADMIN') checks for exactly 'ROLE_ADMIN'. Since user roles are stored as 'ROLE_ADMIN' in the database, I use hasAuthority('ROLE_ADMIN') consistently to avoid confusion.")

add_qa(doc, 15,
       "What happens if a JWT token expires?",
       "The JwtUtil.isTokenExpired() method checks the token's expiration claim against the current time. If expired, validateToken() returns false. The JwtFilter will not set the SecurityContext, so the request is treated as unauthenticated and returns a 401 Unauthorized response. The frontend should detect this and redirect the user to the login page to obtain a new token.")

add_h2(doc, "Section D — React / Frontend Questions")

add_qa(doc, 16,
       "How does protected routing work in your React app?",
       "ProtectedRoute.jsx checks if a JWT token exists in localStorage. If not, it redirects to /login using React Router's <Navigate> component. If the token exists, the protected component is rendered. Role-based routing is handled by decoding the JWT using jwt-decode to extract the role field and rendering admin or user routes accordingly.")

add_qa(doc, 17,
       "Why did you use Vite instead of Create React App?",
       "Vite uses native ES modules and esbuild for bundling, making dev server startup near-instant compared to CRA's Webpack-based approach. Hot Module Replacement (HMR) is much faster. Build times are significantly reduced. Vite also has better TypeScript support and a modern plugin ecosystem. For a project like this, Vite provides a better developer experience.")

add_qa(doc, 18,
       "How did you manage API calls in the frontend?",
       "API base URL is configured pointing to the Railway backend. Axios is used for HTTP calls with the Authorization header automatically set from localStorage token. Each page component uses useEffect to fetch data on mount. Custom fetch calls are also used in some components. Loading states are managed with useState(false) to show loaders. Error handling uses try/catch.")

add_qa(doc, 19,
       "Explain how the loan status notification modal works.",
       "On Dashboard load, the latest loan application is sorted by date and ID. If its status is 'approved' or 'rejected', the component checks localStorage for a key last_notified_{status}_{userId}. If the loan ID doesn't match the stored key, a Modal is shown with an animated icon (CheckCircle with animate-bounce for approved, XCircle for rejected). Once shown, the loan ID is saved to localStorage so the modal doesn't appear again for the same loan. This provides a one-time notification UX.")

add_qa(doc, 20,
       "What is Tailwind CSS and how is it different from Bootstrap?",
       "Tailwind CSS is a utility-first CSS framework — instead of pre-built components (like Bootstrap's buttons, cards), Tailwind provides low-level utility classes (flex, p-8, rounded-2xl, bg-gradient-to-br, etc.) that you compose in JSX. This gives full design control without writing custom CSS. Bootstrap has opinionated default styles that require overriding. Tailwind results in smaller final CSS bundles (purges unused classes) and is more flexible for custom designs.")

add_qa(doc, 21,
       "What is Recharts and how did you use it in the admin dashboard?",
       "Recharts is a composable charting library built on D3.js for React. In the admin dashboard, I used it to display loan statistics — BarChart for total/approved/pending/rejected counts and potentially PieChart for distribution percentages. The data comes from the /api/loans/stats endpoint which returns counts aggregated in the backend.")

add_h2(doc, "Section E — AI / Python / Microservice Questions")

add_qa(doc, 22,
       "Explain the Python AI microservice architecture.",
       "The AI microservice is a Python FastAPI application deployed on Render. It exposes two POST endpoints: /predict (returns binary approval: true/false) and /recommend (returns a ranked list of suitable lenders). At startup, a pre-trained scikit-learn model is loaded from loan_model.pkl using pickle. When a request arrives, the 7 input features (age, income, existing_emi, credit_score, asset_value, loan_amount, tenure) are converted to a NumPy array and passed to model.predict(). Pydantic validates the request schema.")

add_qa(doc, 23,
       "What machine learning model did you train for loan prediction?",
       "A classification model trained on a loan dataset with features: age, income, existing EMI, credit score, asset value, loan amount, and tenure. The model predicts binary output (approved: 1 or 0). The trained model is serialized with pickle and loaded at FastAPI startup. scikit-learn models like RandomForestClassifier or LogisticRegression are typical choices for this kind of binary classification task.")

add_qa(doc, 24,
       "What is FastAPI and why did you choose it over Flask?",
       "FastAPI is an async Python web framework based on Starlette and Pydantic. Compared to Flask: FastAPI has built-in data validation via Pydantic models, automatic OpenAPI/Swagger documentation, native async support, and type hints everywhere. It's significantly faster than Flask for I/O bound tasks. For a microservice exposing simple ML prediction endpoints, FastAPI's speed and auto-validation make it ideal.")

add_h2(doc, "Section F — Database & Design Questions")

add_qa(doc, 25,
       "How did you design the database schema? Why is LoanApplication linked to both User and Asset?",
       "LoanApplication is the central transaction table. It links to User (the borrower), LoanProvider (the bank), and Asset (the collateral). This mirrors real-world loan structure: a bank (LoanProvider) lends money to a person (User) against a collateral (Asset). Using @ManyToOne relationships with foreign keys ensures data integrity. The @JsonBackReference on User in LoanApplication prevents infinite JSON recursion during serialization.")

add_qa(doc, 26,
       "What is spring.jpa.hibernate.ddl-auto=update and what are its risks?",
       "'update' tells Hibernate to compare the entity model with the existing schema and apply non-destructive changes (add columns, add tables). It will NOT drop columns or tables. This is convenient for development but risky in production because: 1) It may silently fail for complex schema changes. 2) It doesn't handle column renames or type changes. For production, Flyway or Liquibase migration tools are recommended.")

add_qa(doc, 27,
       "Why is @JsonIgnore used on the Asset entity's user field?",
       "@JsonIgnore prevents the user field from being included in the JSON response when an Asset is serialized. Without it, serializing an Asset would include the full User object, which would include all their assets, creating an infinite recursive loop. Similarly, @JsonBackReference on the User field in LoanApplication prevents recursive serialization of the bidirectional relationship.")

add_h2(doc, "Section G — Deployment & DevOps Questions")

add_qa(doc, 28,
       "How is the application deployed and what is the CI/CD strategy?",
       "Frontend on Netlify: connected to Git repository; auto-deploys on push. Backend on Railway: Docker-based deployment using the Dockerfile in the backend folder. AI on Render: Python service deployed from the AI/loanmaker-ai folder. Database on Aiven: managed MySQL cloud, connection string in application.properties using environment variable injection (${PORT:8080}, ${JWT_SECRET:...}). Sensitive values (JWT secret, DB password) are configured as environment variables in each platform, not hardcoded.")

add_qa(doc, 29,
       "What does the Dockerfile for the backend look like conceptually?",
       "The Dockerfile packages the Spring Boot application as a JAR using Maven, then runs it in a JRE container. Typically: FROM maven:3.9-eclipse-temurin-21 to build, then FROM eclipse-temurin:21-jre for the runtime image. The JAR is copied and run with java -jar. The PORT environment variable from Railway is passed via server.port=${PORT:8080} in application.properties.")

add_qa(doc, 30,
       "How did you handle CORS between different deployment domains?",
       "CORS is configured in SecurityConfig.corsConfigurationSource() to allow specific origins: http://localhost:5173 for local development and https://timely-tanuki-loan-maker.netlify.app for production. All HTTP methods (GET, POST, PUT, DELETE, OPTIONS) are allowed. Credentials must be allowed (setAllowCredentials(true)) since JWT tokens are sent in headers. The configuration is applied to all routes using source.registerCorsConfiguration('/**', configuration).")

add_h2(doc, "Section H — Scenario / Problem-Solving Questions")

add_qa(doc, 31,
       "What challenges did you face and how did you solve them?",
       "1) Circular JSON serialization: Solved using @JsonBackReference and @JsonIgnore on bidirectional JPA relationships. 2) CORS issues in production: Solved by configuring CorsConfigurationSource in SecurityConfig with specific allowed origins. 3) Spring Security blocking endpoints: Debugged by understanding the difference between hasRole() and hasAuthority() — roles stored as ROLE_ADMIN require hasAuthority(). 4) AI microservice cold starts on Render free tier: Handled gracefully with try/catch in LoanAIService. 5) Token expiry UX: Managed on frontend by checking 401 responses and clearing localStorage.")

add_qa(doc, 32,
       "How would you improve this project in the future?",
       "1) Add Refresh Tokens to extend sessions without re-login. 2) Implement Flyway/Liquibase for proper database migrations. 3) Add Spring Cache (Redis) for caching loan provider data. 4) Implement notification emails using Spring Mail when loan status changes. 5) Add unit tests with JUnit/Mockito for service layer. 6) Use React Query or Redux Toolkit for better server state management. 7) Add rate limiting to the auth endpoints to prevent brute force attacks. 8) Train the ML model on a larger dataset and add model versioning using MLflow.")

add_qa(doc, 33,
       "How would you secure the application better?",
       "1) Move JWT to HTTP-only cookies instead of localStorage to prevent XSS attacks. 2) Add refresh token rotation. 3) Use HTTPS everywhere (already done via cloud deployments). 4) Add @Valid bean validation on all request bodies. 5) Rate limiting on /api/auth/login. 6) Mask sensitive fields (password) in logs. 7) Store DB credentials in AWS Secrets Manager or similar vault. 8) Add API Gateway for centralized auth and rate limiting.")

add_qa(doc, 34,
       "What is the difference between @RestController and @Controller?",
       "@Controller is a traditional MVC controller that returns view names (Thymeleaf templates, JSP). @RestController is a combination of @Controller + @ResponseBody, meaning every method automatically serializes the return value to JSON and writes it directly to the HTTP response. Since this is a REST API with a separate React frontend, all controllers use @RestController.")

add_qa(doc, 35,
       "If two users apply for a loan at the same time, is there a race condition?",
       "With the current design, there's minimal risk since each loan application is an independent INSERT operation. However, for scenarios like checking and updating a shared resource simultaneously (e.g., remaining loan quota of a provider), we would need @Transactional with proper isolation levels. Spring's @Transactional with SERIALIZABLE isolation would prevent phantom reads. JPA's @Version annotation can enable optimistic locking.")

add_h2(doc, "Section I — Conceptual / Technical Theory Questions")

add_qa(doc, 36,
       "What is the difference between authentication and authorization?",
       "Authentication is verifying WHO you are — proving identity using credentials (username/password). In this project, login returns a JWT proving identity. Authorization is verifying WHAT you can do — checking permissions. In this project, @PreAuthorize('hasAuthority(\"ROLE_ADMIN\")') ensures only admins can access admin endpoints, even if a regular user is authenticated.")

add_qa(doc, 37,
       "What is @Transactional and when would you use it here?",
       "@Transactional wraps a method in a database transaction — if any operation fails, all changes are rolled back. For example, if a loan application save fails after updating the user's previousLoans count, @Transactional would roll back both operations, keeping data consistent. Use it on service methods that perform multiple database operations that must succeed or fail together.")

add_qa(doc, 38,
       "What is the purpose of @GeneratedValue(strategy = GenerationType.IDENTITY)?",
       "It tells JPA to use the database's auto-increment feature to generate the primary key value. The database assigns the next available integer ID when a record is inserted. IDENTITY strategy is best for MySQL. Other strategies: SEQUENCE (Oracle/PostgreSQL), TABLE (database-agnostic), AUTO (lets JPA choose).")

add_qa(doc, 39,
       "How does React's useEffect work and how did you use it?",
       "useEffect is a hook that runs side effects (data fetching, event listeners, timers) after the component renders. The dependency array controls when it re-runs: empty [] means run only on mount. [token] means re-run when token changes. I used useEffect to fetch user profile and loan data on Dashboard mount, and to add/remove event listeners for the eligibility-updated custom event and browser visibility changes for real-time data refresh.")

add_qa(doc, 40,
       "What is Pydantic and how does it help in the AI microservice?",
       "Pydantic is a Python data validation library. In FastAPI, request bodies are declared as Pydantic BaseModel subclasses. FastAPI automatically validates incoming JSON against the model, returning a 422 Unprocessable Entity error if types are wrong or required fields are missing. LoanRequest model ensures all 7 fields (age: int, income: float, credit_score: int, etc.) are present and correctly typed before the ML model runs.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. PROJECT SUMMARY CARD
# ═══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "8. PROJECT AT A GLANCE (Quick Reference)")

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Details'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

rows_data = [
    ("Project Name", "Loan Maker — AI-Powered Loan Management System"),
    ("Backend", "Spring Boot 3.5.5, Java 21, Maven"),
    ("Frontend", "React 19, Vite 7, Tailwind CSS v4"),
    ("Database", "MySQL 8 on Aiven Cloud"),
    ("Authentication", "JWT (JJWT 0.11.5) + BCrypt + Spring Security"),
    ("AI Microservice", "Python FastAPI + scikit-learn + NumPy"),
    ("Deployment", "Netlify (FE) | Railway (BE) | Render (AI)"),
    ("Key Libraries", "Axios, Recharts, Framer Motion, Lucide, react-router-dom v7"),
    ("API Endpoints", "Auth, Loans (CRUD), Assets, Admin, AI (predict/recommend)"),
    ("User Roles", "USER (apply loans, manage assets) | ROLE_ADMIN (all + manage users)"),
    ("Loan Statuses", "PENDING → APPROVED or REJECTED"),
    ("AI Features", "Rule-based evaluation + ML model prediction + lender recommendation"),
]
for cat, detail in rows_data:
    row = table.add_row().cells
    row[0].text = cat
    row[1].text = detail

doc.add_paragraph()
add_body(doc, "This document was auto-generated from the Loan Maker project source code analysis.")

# Save – two locations for easy access
import os
paths = [
    r"c:\Users\pechi\OneDrive\Desktop\Loan_Maker_Project\Loan_Maker_Project_Interview_Guide.docx",
    r"c:\Users\pechi\Desktop\Loan_Maker_Project_Interview_Guide.docx",
    r"c:\Users\pechi\OneDrive\Desktop\Loan_Maker_Project_Interview_Guide.docx",
]
saved = []
for p in paths:
    try:
        doc.save(p)
        print(f"Saved: {p}")
        saved.append(p)
    except Exception as e:
        print(f"Could not save to {p}: {e}")
if saved:
    print(f"\nDocument successfully saved at:\n" + "\n".join(saved))
else:
    print("ERROR: Could not save document anywhere!")
